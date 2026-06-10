from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=18, color=(0, 51, 102))
    elif level == 2:
        set_font(run, bold=True, size=14, color=(0, 82, 155))
    else:
        set_font(run, bold=True, size=12, color=(30, 30, 30))
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=bold, size=10.5, italic=italic, color=color)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(180, 0, 0)
    # light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F5F5F5')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, header_color='003366'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], header_color)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9.5)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'EEF4FF')
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('BizOS')
set_font(run, bold=True, size=36, color=(0, 51, 102))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mobile API Integration Guide')
set_font(run, bold=False, size=20, color=(0, 82, 155))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for Mobile Developer  |  Version 1.0  |  2026')
set_font(run, italic=True, size=11, color=(100, 100, 100))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Overview
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Overview')
add_body(doc, 'BizOS is a business operations management platform. The mobile app shares the same backend as the web app. Any data updated in the web app is immediately reflected in the mobile app and vice versa, because both use the same REST API.')
add_body(doc, 'Base URL (Production):', bold=True)
add_code(doc, 'https://bizos.dio.lk/panda_jp/backend/public/api')
add_body(doc, 'All endpoints are prefixed with /api. Example full URL:')
add_code(doc, 'https://bizos.dio.lk/panda_jp/backend/public/api/login')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Authentication
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Authentication')
add_body(doc, 'BizOS uses Laravel Sanctum Bearer Token authentication.')
add_body(doc, 'How it works:', bold=True)
add_body(doc, '1. Call POST /api/login with email and password.')
add_body(doc, '2. Receive a token in the response.')
add_body(doc, '3. Include the token in ALL subsequent protected requests as an HTTP header:')
add_code(doc, 'Authorization: Bearer <your_token>')
add_body(doc, 'Also set the following headers on all requests:', bold=True)
add_code(doc, 'Accept: application/json\nContent-Type: application/json')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Auth Endpoints
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Authentication Endpoints')

add_heading(doc, '3.1  Login', level=2)
add_body(doc, 'POST  /api/login', bold=True)
add_body(doc, 'No authentication required.')
add_code(doc, '// Request body (JSON)\n{\n  "email": "user@example.com",\n  "password": "password123"\n}')
add_code(doc, '// Response\n{\n  "token": "1|abc123xyz...",\n  "user": { "id": 1, "name": "John", "email": "...", ... }\n}')
add_body(doc, 'Save the token securely (e.g., device keychain). Use it for all future requests.', italic=True, color=(80,80,80))

add_heading(doc, '3.2  Verify 2FA (if enabled)', level=2)
add_body(doc, 'POST  /api/verify-2fa', bold=True)
add_code(doc, '{\n  "user_id": 1,\n  "otp": "123456"\n}')

add_heading(doc, '3.3  Logout', level=2)
add_body(doc, 'POST  /api/logout', bold=True)
add_body(doc, 'Requires: Authorization: Bearer <token>')
add_body(doc, 'Invalidates the current token on the server.')

add_heading(doc, '3.4  Register New User', level=2)
add_body(doc, 'POST  /api/register', bold=True)
add_code(doc, '{\n  "name": "John Doe",\n  "email": "john@example.com",\n  "password": "password",\n  "password_confirmation": "password",\n  "mobile": "0771234567",\n  "isCompanyEmployee": true,\n  "department": "HR",\n  "jobPosition": "Manager",\n  "employeeNumber": "EMP001"\n}')

add_heading(doc, '3.5  Forgot Password Flow', level=2)
add_body(doc, 'Step 1 — Send OTP to email:', bold=True)
add_code(doc, 'POST /api/forgot-password\n{ "email": "user@example.com" }')
add_body(doc, 'Step 2 — Verify OTP:', bold=True)
add_code(doc, 'POST /api/reset-password\n{ "email": "user@example.com", "otp": "123456" }')
add_body(doc, 'Step 3 — Set new password:', bold=True)
add_code(doc, 'POST /api/change-password\n{ "email": "user@example.com", "password": "newpassword" }')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Current User / Profile
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Current User & Profile')

add_heading(doc, '4.1  Get Logged-in User', level=2)
add_body(doc, 'GET  /api/user', bold=True)
add_body(doc, 'Returns the full user object: name, email, role, permissions, profile image URL, preferences, etc.')
add_body(doc, 'Use this on app launch to check who is logged in and what they can access.')

add_heading(doc, '4.2  Update Profile Details', level=2)
add_body(doc, 'POST  /api/user/{id}/profile-update', bold=True)
add_code(doc, '{\n  "name": "New Name",\n  "mobile": "0771234567",\n  "gender": "Male",\n  "themePreference": "dark",\n  "defaultLandingPage": "/home"\n}')

add_heading(doc, '4.3  Upload Profile Image', level=2)
add_body(doc, 'POST  /api/user/{id}/profile-update', bold=True)
add_body(doc, 'Use multipart/form-data. Field name: profileImage[0]')
add_code(doc, 'Content-Type: multipart/form-data\nField: profileImage[0] = <image file>')

add_heading(doc, '4.4  Remove Profile Image', level=2)
add_body(doc, 'POST  /api/user/{id}/profile-update', bold=True)
add_code(doc, '{ "clearImage": true }')

add_heading(doc, '4.5  Change Password (while logged in)', level=2)
add_body(doc, 'POST  /api/user-change-password', bold=True)
add_code(doc, '{\n  "currentPassword": "oldpassword",\n  "newPassword": "newpassword",\n  "newPassword_confirmation": "newpassword"\n}')

add_heading(doc, '4.6  Change Email (3-step)', level=2)
add_code(doc, 'Step 1:  POST /api/user/{id}/email-change\n         { "currentEmail": "old@email.com" }\n\nStep 2:  POST /api/user/{id}/email-change-verify\n         { "otp": "123456" }\n\nStep 3:  POST /api/user/{id}/email-change-confirm\n         { "newEmail": "new@email.com" }')

add_heading(doc, '4.7  Active Sessions', level=2)
add_code(doc, 'GET    /api/user/active-sessions          — list all active sessions\nDELETE /api/user/active-sessions/{tokenId} — revoke a session')

add_heading(doc, '4.8  Login History', level=2)
add_code(doc, 'GET  /api/user/login-history')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Users (Admin)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Users (Admin)')
add_body(doc, 'These endpoints manage all users in the system. Required permissions: ADMIN_USERS_VIEW and ADMIN_USERS_EDIT.')

add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET',    '/api/all-users',          'None',           'List all users (public)'],
        ['GET',    '/api/users',              'Bearer Token',   'Admin full user list'],
        ['GET',    '/api/users/search?q=',    'None',           'Search users by name/email'],
        ['POST',   '/api/users/{id}/update',  'Bearer Token',   'Update user role/department/status'],
        ['DELETE', '/api/users/{id}',         'Bearer Token',   'Delete user & revoke all tokens'],
        ['GET',    '/api/users-assignee-level','Bearer Token',  'Get assignee levels list'],
    ]
)

add_heading(doc, 'Update User — Request Body', level=3)
add_code(doc, '{\n  "name": "Updated Name",\n  "userType": "3",\n  "department": "Finance",\n  "assigneeLevel": "1",\n  "jobPosition": "Manager",\n  "availability": true\n}')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Dashboard
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Dashboard')
add_body(doc, 'GET  /api/dashboard', bold=True)
add_body(doc, 'Required permission: DASHBOARD_VIEW')
add_body(doc, 'Returns all configured systems with their latest KPI values.')
add_code(doc, '// Response structure\n{\n  "systems": [\n    {\n      "id": 1,\n      "name": "Finance System",\n      "type": "finance",\n      "metrics": [\n        { "id": 1, "label": "TOTAL DEBTORS", "value": "1,330,000", "note": null }\n      ],\n      "badge": { "label": "ABSENTEEISM", "value": "4.2%", "type": "danger" }\n    },\n    {\n      "id": 2,\n      "name": "Stores",\n      "type": "stores",\n      "metrics": [...],\n      "alert": { "label": "SPARE PART UNAVAILABILITY", "value": "200 min" },\n      "status": "System Optimal"\n    }\n  ]\n}')

add_body(doc, 'System type values:', bold=True)
add_table(doc,
    ['Type', 'Description'],
    [
        ['finance',     'Finance system — shows metrics grid + absenteeism badge'],
        ['stores',      'Stores system — shows metrics grid + unavailability alert'],
        ['healthSafety','Health & Safety — shows accidents badge + metrics'],
        ['maintenance', 'Maintenance — shows metrics grid'],
        ['default',     'Custom/other systems — shows metrics grid'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Systems & KPIs
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Systems & KPIs')

add_table(doc,
    ['Method', 'Endpoint', 'Permission', 'Description'],
    [
        ['GET',    '/api/systems',                         'Auth',              'List all systems'],
        ['GET',    '/api/systems/{id}',                    'Auth',              'Get single system'],
        ['POST',   '/api/systems',                         'SYSTEM_SETUP_EDIT', 'Create system'],
        ['PUT',    '/api/systems/{id}',                    'SYSTEM_SETUP_EDIT', 'Update system'],
        ['DELETE', '/api/systems/{id}',                    'SYSTEM_SETUP_EDIT', 'Delete system'],
        ['GET',    '/api/systems/{id}/metrics',            'Auth',              'Get KPIs for a system'],
        ['POST',   '/api/systems/{id}/metrics',            'SYSTEM_SETUP_EDIT', 'Add KPI to system'],
        ['PUT',    '/api/metrics/{id}',                    'SYSTEM_SETUP_EDIT or INPUT_PAGE_EDIT', 'Update KPI value'],
        ['DELETE', '/api/metrics/{id}',                    'SYSTEM_SETUP_EDIT', 'Delete a KPI'],
        ['GET',    '/api/metrics/{id}/daily-values',       'INPUT_PAGE_VIEW or DASHBOARD_VIEW', 'Get KPI history'],
        ['GET',    '/api/metrics/{id}/daily-values/{date}','INPUT_PAGE_VIEW or DASHBOARD_VIEW', 'Get KPI value for date'],
        ['POST',   '/api/metrics/{id}/daily-values',       'INPUT_PAGE_EDIT',   'Add KPI daily entry'],
    ]
)

add_body(doc, 'Date format for {date}: YYYY-MM-DD', italic=True, color=(80,80,80))

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — Company Settings
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Company Settings')

add_table(doc,
    ['Method', 'Endpoint', 'Permission', 'Description'],
    [
        ['GET',    '/api/organizations',              'None',                'Get company/org info'],
        ['POST',   '/api/organizations/{id}/update',  'COMPANY_SETTING_EDIT','Update company info'],
        ['DELETE', '/api/organizations/{id}/delete',  'COMPANY_SETTING_EDIT','Delete org record'],
        ['GET',    '/api/company-holidays',           'Auth',               'List holidays'],
        ['GET',    '/api/company-holidays/{id}',      'Auth',               'Get single holiday'],
        ['POST',   '/api/company-holidays',           'COMPANY_SETTING_EDIT','Add holiday'],
        ['PUT',    '/api/company-holidays/{id}',      'COMPANY_SETTING_EDIT','Update holiday'],
        ['DELETE', '/api/company-holidays/{id}',      'COMPANY_SETTING_EDIT','Delete holiday'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — Access Management
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Access Management')
add_body(doc, 'Manage user permission roles (who can see/edit what in the system).')

add_table(doc,
    ['Method', 'Endpoint', 'Permission', 'Description'],
    [
        ['GET',    '/api/user-permissions',              'ADMIN_ACCESS_MNG_VIEW',  'List all roles'],
        ['GET',    '/api/user-permissions/{id}/show',    'ADMIN_ACCESS_MNG_VIEW',  'Get single role'],
        ['POST',   '/api/user-permissions',              'ADMIN_ACCESS_MNG_CREATE','Create new role'],
        ['POST',   '/api/user-permissions/{id}/update',  'ADMIN_ACCESS_MNG_EDIT',  'Update role permissions'],
        ['DELETE', '/api/user-permissions/{id}/delete',  'ADMIN_ACCESS_MNG_DELETE','Delete role'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Reference / Lookup Data
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Reference / Lookup Data')
add_body(doc, 'These endpoints return dropdown/lookup values. Most do not require authentication.')

add_table(doc,
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',  '/api/departments',         'List departments'],
        ['POST', '/api/departments',         'Add department'],
        ['GET',  '/api/job-positions',       'List job positions'],
        ['POST', '/api/job-positions',       'Add job position'],
        ['GET',  '/api/assignee-level',      'List assignee levels'],
        ['GET',  '/api/responsible-section', 'List responsible sections'],
        ['POST', '/api/responsible-section', 'Add responsible section'],
        ['GET',  '/api/factory',             'Get factory info'],
        ['POST', '/api/factory',             'Add factory'],
        ['GET',  '/api/user-types',          'List user types'],
        ['GET',  '/api/person-types',        'List person types'],
        ['GET',  '/api/hr-divisions',        'List HR divisions'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — Image Upload
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '11. Image Upload (General)')

add_table(doc,
    ['Method', 'Endpoint', 'Description'],
    [
        ['POST',   '/api/upload',                  'Upload an image (multipart/form-data)'],
        ['GET',    '/api/image/{imageId}',          'Retrieve image by ID'],
        ['DELETE', '/api/image/{imageId}',          'Delete image'],
        ['POST',   '/api/image/update/{imageId}',   'Replace an existing image'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — Permission Keys
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '12. Permission Keys Reference')
add_body(doc, 'After login, the user object contains a permissionObject field. Check these keys (true/false) to show or hide features in the mobile app.')

add_table(doc,
    ['Permission Key', 'What It Controls'],
    [
        ['DASHBOARD_VIEW',          'View the main KPI dashboard'],
        ['SYSTEM_SETUP_VIEW',       'View system/KPI setup'],
        ['SYSTEM_SETUP_EDIT',       'Create/edit/delete systems and KPIs'],
        ['INPUT_PAGE_VIEW',         'View KPI data entry page'],
        ['INPUT_PAGE_EDIT',         'Enter KPI daily values'],
        ['ADMIN_USERS_VIEW',        'View users list'],
        ['ADMIN_USERS_EDIT',        'Edit or delete users'],
        ['ADMIN_ACCESS_MNG_VIEW',   'View access roles'],
        ['ADMIN_ACCESS_MNG_EDIT',   'Edit access roles'],
        ['ADMIN_ACCESS_MNG_CREATE', 'Create new access roles'],
        ['ADMIN_ACCESS_MNG_DELETE', 'Delete access roles'],
        ['COMPANY_SETTING_EDIT',    'Edit company/org settings'],
        ['USER_SETTING_VIEW',       'View own profile settings'],
        ['USER_SETTING_EDIT',       'Edit own profile, password, email'],
        ['CUSTOM_PAGE_VIEW',        'View custom dashboard page'],
        ['CUSTOM_PAGE_EDIT',        'Edit custom dashboard page'],
    ]
)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — Error Responses
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '13. Error Responses')

add_table(doc,
    ['HTTP Status', 'Meaning', 'Action'],
    [
        ['200', 'OK',                    'Request succeeded'],
        ['401', 'Unauthenticated',       'Token missing, expired, or revoked — redirect to login'],
        ['403', 'Forbidden',             'User lacks required permission — hide the feature'],
        ['404', 'Not Found',             'Resource does not exist'],
        ['422', 'Validation Error',      'Check the "errors" field in the JSON response body'],
        ['500', 'Server Error',          'Backend issue — show a generic error message'],
    ]
)

add_body(doc, 'Validation error response shape:', bold=True)
add_code(doc, '{\n  "message": "The email field is required.",\n  "errors": {\n    "email": ["The email field is required."]\n  }\n}')

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — Quick Start Checklist
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '14. Quick Start Checklist for Mobile Developer')

steps = [
    'Set base URL: https://bizos.dio.lk/panda_jp/backend/public/api',
    'Call POST /api/login to get token',
    'Store token securely (device keychain recommended)',
    'Add Authorization: Bearer <token> header to all protected requests',
    'Call GET /api/user to get current user info and permissionObject',
    'Use permissionObject keys to show/hide features in the UI',
    'Call GET /api/dashboard to display KPI overview',
    'On logout, call POST /api/logout and clear the stored token',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(step)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
add_body(doc, 'For questions about this API, contact the backend team at: skytech.srilanka@gmail.com', italic=True, color=(80,80,80))

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'e:/bisiness_OS-main/BizOS_API_Documentation.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
